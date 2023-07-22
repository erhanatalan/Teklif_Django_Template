# -*- coding: utf-8 -*-
import docx
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from datetime import date
from .get_info_from_goolesheets import *
# import time
import os
import locale
from .veri import *
if(usmodel=='B'):
    from .infoB import *
elif(usmodel=='C'):
    from .infoC import *
elif(usmodel=='CB'):
    from .infoCB import *
elif(usmodel=='CC'):
    from .infoCC import *
elif(usmodel=='V'):
    from .infoV import *
elif(usmodel=='I'):
    from .infoI import *



def toword():
    # ******************************************
    # Türkçe lokal ayarlarını kullan
    default_locale = locale.getdefaultlocale()
    locale.setlocale(locale.LC_TIME, default_locale)
    # Yeni bir Word belgesi oluştur
    doc = docx.Document()
    # ******************************************
    
    # Font ailesini değiştirmek için stil şablonu oluştur
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arimo'  # Yeni font ailesi
    font.size = Pt(10)  # Yeni yazı boyutu
    # ******************************************
    # Bugünün tarihini al
    bugun = date.today()
    gun = bugun.strftime("%d")
    ay = bugun.strftime("%B")
    yil = bugun.strftime("%Y")
    # ******************************************
    # Belge kenar boşluklarını ayarla
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)      # Sol kenar boşluğunu 1 inç olarak ayarla
        section.right_margin = Inches(0.5)     # Sağ kenar boşluğunu 1 inç olarak ayarla
        section.top_margin = Inches(0.1)       # Üst kenar boşluğunu 1 inç olarak ayarla
        section.bottom_margin = Inches(0.1)    # Alt kenar boşluğunu 1 inç olarak ayarla
    # ******************************************
    # Paragraf oluştur
    paragraf = doc.add_paragraph()
    paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ******************************************
    # Resmi ekle
    resim_dizini = 'media/'
    resim_yolu = os.path.join(resim_dizini, 'abs.png')
    paragraf.add_run().add_picture(resim_yolu)
    # Sayfa genişliğini al
    section = doc.sections[0]
    sayfa_genisligi = section.page_width - section.left_margin - section.right_margin
    # Resmi sayfa genişliği kadar ayarla
    resim_genislik = sayfa_genisligi - Inches(0.5)  # Kenar boşluklarını da düşünmek için biraz azaltma yapabilirsiniz
    paragraf.width = resim_genislik
    # ******************************************
    # Tarihi belgeye ekle ve en sağa yasla
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f'{gun} {ay} {yil}')
    run.font.size = Pt(12)  # Yazı boyutunu ayarla
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Hizalamayı sağa ayarla
    # ******************************************
    # Başlık ekle
    heading = doc.add_heading('KANTAR FİYAT TEKLİFİ', level=3)
    # Başlık metnini siyah ve kalın yap
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk (RGB: 0, 0, 0)
    run.bold = True
    run.font.size = Pt(18)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ******************************************
    # Metin paragrafları ekle
    satir1 = doc.add_paragraph(f'KONU: {adet} Adet 3X{uzunluk}m, Tam Elektronik Zemin Üstü {tonaj} Ton Kapasiteli {model} Model {mekanik} Kantar İmalat ve Kurulum Fiyat Teklifi.')

    satir2 = doc.add_paragraph(f'1.) MEKANİK VE ELEKTRONİK AKSAM: Teknik bilgiler, İnşaat Projesi ve Fotoğraflar ikinci sayfada yer almaktadır.')

    satir3 = doc.add_paragraph(f'2.) İNŞAAT VE PROJE: Kantar zemini için proje verilecektir. Proje üzerindeki detaylar, inşaat, kabin, hafriyat ve çevre düzenlemesi {insaat} yapılacaktır.')

    # Metin ve linki tanımla
    # metin = f'{linkB6}'
    # satir4 = doc.add_paragraph('İnşaat Projesi :')
    # satir41 = add_hyperlink(satir4, 'https://drive.google.com/file/d/1el2eAMfcprVtToEwUxMSUF60ZRvoiHOm/view?usp=sharing', 'B6 inşaat projesi.pdf', '1E90FF', False)
    # run = satir4.add_run(satir41)
    # Metin ve linki tanımla

    satir6 = doc.add_paragraph(f'3) TESLİMAT: Anlaşma; sözleşmenin yapıldığı tarihten (ön peşinat ödemesi yapıldığı) tarihten itibaren {teslimatgun} iş günü içerisinde kantar teslim edilecektir.')

    satir7 = doc.add_paragraph(f'4) MONTAJ: Kantar firmamız elemanları tarafından kurulacaktır. Gerekli eğitim kullanıcıya verilecektir. Kantarın kurulumu sırasında gerekli olan vinç (3-4 saat) vb. ekipmanlar {vinc} tedarik edilecektir.')

    satir8 = doc.add_paragraph(f'5) NAKLİYE: Kantarın nakliyesi {nakliye} yapılacaktır. Kantar platformları toplam {rampaharic} Ton {dahil} gelecektir. Nakliyesi için bir tır yeterlidir.')

    satir9 = doc.add_paragraph(f'6) FİYAT: Fiyatlarda %{kdv} KDV dâhil değildir.\n1 Adet fiyatı {fiyat}₺ {rmp}\n1 Adet Periyodik Muayene Sertifikası Ücretsiz (2 Yıl Geçerli)')

    satir10 = doc.add_paragraph(f'7) ÖDEME:\n\ta.) %{onodeme} Peşin siparişte, geri kalan tutar teslimatta nakit olarak ödenir.\n\tb.) Diğer bir ödeme planı %{onodeme} nakit siparişte, kalanı teslimatta {cekvade} gün vade çek olarak ödenirse fiyat üzerine %{ekcekodeme} eklenir.')

    satir11 = doc.add_paragraph(f'8) GARANTİ VE TEKLİF SÜRESİ: Elektronik Aksam 2 yıl, Mekanik Aksam {mekanikgaranti} yıl Garantilidir. Teklifimiz {teklifsuresi} gün geçerlidir.')

    satir12 = doc.add_paragraph(f'Ülkemizin tüm illerinde ayrıca ARABİSTAN, IRAK, TÜRKMENİSTAN, BULGARİSTAN, RUSYA, GÜRCİSTAN VE BİRÇOK AFRİKA ÜLKESİNDE’ da referanslarımız, ayrıca Servis hizmetimiz vardır.')

    satir13 = doc.add_paragraph(f'Teklifimizi uygun karşılayacağınızı umar, işlerinizde başarılar dileriz.Saygılarımla…')

    satir14 = doc.add_paragraph(f'Saygılarımla…')
    # ******************************************
    # Tabloyu oluştur
    rows = 3
    cols = 2
    table = doc.add_table(rows=rows, cols=cols)
    # Tablonun kenarlıklarını etkinleştir
    table.style = 'Table Grid'
    # Tablonun hücrelerine veri ekle
    table.cell(0, 0).text = 'ALICI FİRMA'
    table.cell(0, 1).text = 'SATICI FİRMA'
    table.cell(1, 0).text = f'{firma}'
    table.cell(1, 1).text = f'ABS ELEKTRONİK TARTI SİSTEMLERİ SAN.TİC.LTD.ŞTİ.'
    table.cell(2, 0).text = f'KAŞE İMZA\n\n\n\n\n'
    cell = table.cell(2, 1)  # İstenilen hücrenin konumunu belirleyin
    cell_paragraph = cell.paragraphs[0]  # Hücredeki ilk paragrafı alın
    run = cell_paragraph.add_run()
    resim_yolu = os.path.join(resim_dizini, 'abs1.png')
    run.add_picture(resim_yolu,width=Inches(1.5), height=Inches(1.5) )
    # Tablonun her hücresinin etrafına çerçeve ekle
    for row in table.rows:
        for cell in row.cells:
            borders = cell._element.xpath('.//w:tcBorders')
            for border in borders:
                border.attrib.clear()
                border.attrib[nsdecls('w') + 'val'] = 'single'
                border.attrib[nsdecls('w') + 'color'] = '000000'
                border.attrib[nsdecls('w') + 'sz'] = '2'
                # Çerçeve rengini siyah yap
                border.attrib[nsdecls('w') + 'color'] = RGBColor(0, 0, 0)
    # Tablonun hücrelerini hizala
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ***************AYAR**********************
    # hizalama
    satir1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # satir4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # satir5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir6.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir7.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir8.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir9.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    satir10.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    satir11.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir12.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir13.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    satir14.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    # Belgeyi kaydet
    word_dizini = 'C:/Users/yenic/Desktop/Teklif_Django_Template/otomatik/word/'
    doc.save(f'{word_dizini}/{dosya1}.docx')
    print(f'{dosya1} word belgesi olusturuldu')
    # time.sleep(2)



# toword()








