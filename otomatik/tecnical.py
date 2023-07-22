# -*- coding: utf-8 -*-
from .veri1 import *
from .veri import *
import docx
from .hyperlink import add_hyperlink
from docx.shared import Inches, Pt
from docx.shared import RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
if(usmodel=='B'):
    from .infoB import *
if(usmodel=='C'):
    from .infoC import *
if(usmodel=='CB'):
    from .infoCB import *
if(usmodel=='CC'):
    from .infoCC import *
if(usmodel=='V'):
    from .infoV import *
if(usmodel=='I'):
    from .infoI import *
from .veri import model

def towordtecnical():
    # ******************************************
    # Yeni bir Word belgesi oluştur
    doc = docx.Document()
    # ******************************************
    # *******************************************
    # Font ailesini değiştirmek için stil şablonu oluştur
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arimo'  # Yeni font ailesi
    font.size = Pt(10)  # Yeni yazı boyutu
    # ******************************************
    # Belge kenar boşluklarını ayarla
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)      # Sol kenar boşluğunu 1 inç olarak ayarla
        section.right_margin = Inches(0.5)     # Sağ kenar boşluğunu 1 inç olarak ayarla
        section.top_margin = Inches(0.1)       # Üst kenar boşluğunu 1 inç olarak ayarla
        section.bottom_margin = Inches(0.1)    # Alt kenar boşluğunu 1 inç olarak ayarla
    # ******************************************
    # Paragraf oluştur
    paragraf = doc.add_paragraph()
    paragraf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # Resmi ekle
    resim_dizini = 'C:/Users/yenic/Desktop/Teklif_Django_Template/media/'
    resim_yolu = os.path.join(resim_dizini, 'abs.png')
    paragraf.add_run().add_picture(resim_yolu)
    # Sayfa genişliğini al
    section = doc.sections[0]
    sayfa_genisligi = section.page_width - section.left_margin - section.right_margin
    # Resmi sayfa genişliği kadar ayarla
    resim_genislik = sayfa_genisligi - Inches(0.5)  # Kenar boşluklarını da düşünmek için biraz azaltma yapabilirsiniz
    paragraf.width = resim_genislik
    # ******************************************
    # Başlık ekle
    heading = doc.add_heading(f'{model} KANTAR TEKNİK ÖZELLİKLERİ', level=3)
    # Başlık metnini siyah ve kalın yap
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0, 0, 0)  # Siyah renk (RGB: 0, 0, 0)
    run.bold = True
    run.font.size = Pt(18)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # ******************************************
    baslik_metni1 = "PROJELER:" # Başlık metnini ekleyin
    baslik_paragraf1 = doc.add_paragraph() # Başlık paragrafını oluşturun
    baslik_paragraf1.add_run(baslik_metni1).bold = True # Başlık metnini ekleyin ve kalın yapın
    baslik_paragraf1.style = "Heading 2" # Başlık paragrafının stilini ayarlayın
    # ******************************************
     # Metin paragrafları ekle

    
    satir1 = doc.add_paragraph(f'3X{uzunluk}m {model} İnşaat Projesi:')
    # Call the add_hyperlink function to add a hyperlink to the paragraph
    url = f"{proje}"
    text = f"{model}_İnsaat_Projesi"
    underline = True
    add_hyperlink(satir1, url, text, "#0000FF", underline)
    
    if(usmodel=='B'):
        satir11 = doc.add_paragraph(f'3X{uzunluk}m {model} Hazır Rampalı İnşaat Projesi:')
        # Call the add_hyperlink function to add a hyperlink to the paragraph
        url = f"{proje2}"
        text = f"{model}_İnsaat_Projesi"
        underline = True
        add_hyperlink(satir11, url, text, "#0000FF", underline)

    satir2 = doc.add_paragraph(f'{model} Video-Fotoğraflar :')
    # Call the add_hyperlink function to add a hyperlink to the paragraph
    url = f"{foto}"
    text = "ABS_B_Model_Gorsel"
    underline = True
    add_hyperlink(satir2, url, text, "#0000FF", underline)


    baslik_metni5 = "ELEKTRONİK AKSAM MALZEMELERİ:" # Başlık metnini ekleyin
    baslik_paragraf5 = doc.add_paragraph() # Başlık paragrafını oluşturun
    baslik_paragraf5.add_run(baslik_metni5).bold = True # Başlık metnini ekleyin ve kalın yapın
    baslik_paragraf5.style = "Heading 3" # Başlık paragrafının stilini ayarlayın
    # ******************************************
    # Tabloyu oluştur
    table = doc.add_table(rows=9, cols=3)
    table.alignment = WD_ALIGN_VERTICAL.CENTER

    
    # Tablonun hücrelerini dolaşarak genişlikleri ayarla
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx == 0:
                cell.width = docx.shared.Inches(1.5)
            elif idx == 1:
                cell.width = docx.shared.Inches(0.5)
            elif idx == 2:
                cell.width = docx.shared.Inches(2.8)
    # Tablonun hücrelerini dolaşarak hizalamayı ayarla
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
    # Tablonun kenarlıklarını etkinleştir
    table.style = 'Table Grid'
    # Tablonun hücrelerine veri ekle
    
    table.cell(1, 0).text = 'LOADCELL'
    table.cell(1, 1).text = f'{loadcell}'
    table.cell(1, 2).text = 'Ip 68- RC3- 30 TON Keli'
    table.cell(2, 0).text = 'BİLGİSAYAR'
    table.cell(2, 1).text = f'{adet}'
    table.cell(2, 2).text = 'STANDART'
    table.cell(3, 0).text = 'TARTI TERMİNALİ'
    table.cell(3, 1).text = f'{adet}'
    table.cell(3, 2).text = f'{indikator}\nTaksimat:{taksimat}\nHassasiyet:{hassasiyet}'
    table.cell(4, 0).text = 'J-BOX'
    table.cell(4, 1).text = f'{adet}'
    table.cell(4, 2).text = 'ÇOKLU TİP'
    table.cell(5, 0).text = 'BAĞLANTI KUTUSU'
    table.cell(5, 1).text = f'{adet}'
    table.cell(5, 2).text = 'ABS-IP67'
    table.cell(6, 0).text = 'PRİNTER'
    table.cell(6, 1).text = f'{adet}'
    table.cell(6, 2).text = 'JOLİMARK-DP100 & PANTUM-P2500w'
    table.cell(7, 0).text = 'MONİTÖR'
    table.cell(7, 1).text = f'{adet}'
    table.cell(7, 2).text = '17 Inches – 21 Inches'
    table.cell(8, 0).text = 'KLAVYE & MOUSE' 
    table.cell(8, 1).text = f'{adet}'
    table.cell(8, 2).text = 'STANDART'

    # İlk satırdaki metinleri yazdır ve rengini ayarla (kırmızı)
    first_row_cells = table.rows[0].cells
    first_row_cells[0].text = "ÜRÜN"
    first_row_cells[1].text = "ADET"
    first_row_cells[2].text = "TİPİ"
    for i, cell in enumerate(first_row_cells):
        # cell.text = f'{(i+1)}'
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Kırmızı renk

    # ******************************************
    baslik_metni6 = "MEKANİK AKSAM:" # Başlık metnini ekleyin
    baslik_paragraf6 = doc.add_paragraph() # Başlık paragrafını oluşturun
    baslik_paragraf6.add_run(baslik_metni6).bold = True # Başlık metnini ekleyin ve kalın yapın
    baslik_paragraf6.style = "Heading 3" # Başlık paragrafının stilini ayarlayın
    # ******************************************
    # Tabloyu oluştur
    if(usmodel=='B'):
        table = doc.add_table(rows=16, cols=4)
    elif(usmodel=='C'):
        table = doc.add_table(rows=13, cols=4)
    
    
        #******************************


    table.alignment = WD_ALIGN_VERTICAL.CENTER
     # Tablonun kenarlıklarını etkinleştir
    table.style = 'Table Grid'
    # Tablonun hücrelerini dolaşarak genişlikleri ayarla
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            # cell.paragraphs[0].alignment = WD_ALIGN_VERTICAL.CENTER
            if idx == 0:
                cell.width = docx.shared.Inches(2.2)
            elif idx == 1:
                cell.width = docx.shared.Inches(0.8)
            elif idx == 2:
                cell.width = docx.shared.Inches(1.2)
            elif idx == 3:
                cell.width = docx.shared.Inches(1.2)

     # İlk satırdaki metinleri yazdır ve rengini ayarla (kırmızı)
    bir_row_cells = table.rows[0].cells
    bir_row_cells[0].text = "MALZEME"
    bir_row_cells[1].text = "ADET"
    bir_row_cells[2].text = "KALINLIK (mm)"
    bir_row_cells[3].text = "UZUNLUK (mm)"
    for i, cell in enumerate(bir_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = kirmizi 
     # İkinci satırdaki metinleri yazdır 
    iki_row_cells = table.rows[1].cells
    iki_row_cells[0].text = 'ANA TAŞIYICI IPE ÇELİK'
    iki_row_cells[1].text = f'{ipeadet1}'
    iki_row_cells[2].text = f'{kalinlik1}'
    iki_row_cells[3].text = f'{uzunluk1}'
    for i, cell in enumerate(iki_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # Ucuncu satırdaki metinleri yazdır 
    uc_row_cells = table.rows[2].cells
    uc_row_cells[0].text = 'ARA BAĞLANTI IPE ÇELİK'
    uc_row_cells[1].text = f'{adet2}'
    uc_row_cells[2].text = f'{kalinlik2}'
    uc_row_cells[3].text = f'{uzunluk2}'
    for i, cell in enumerate(uc_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # Dortuncu satırdaki metinleri yazdır 
    dort_row_cells = table.rows[3].cells
    dort_row_cells[0].text = f'ARA BAĞLANTI {ara}'
    dort_row_cells[1].text = f'{adet3}'
    dort_row_cells[2].text = f'{kalinlik3}'
    dort_row_cells[3].text = f'{uzunluk3}'
    for i, cell in enumerate(dort_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # Bes satırdaki metinleri yazdır 
    bes_row_cells = table.rows[4].cells
    bes_row_cells[0].text = f'ÜST SAC'
    bes_row_cells[1].text = f'{adet4}m2'
    bes_row_cells[2].text = f'{sackalinligi}mm'
    bes_row_cells[3].text = f'{uzunluk*1000}'
    for i, cell in enumerate(bes_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # Alti satırdaki metinleri yazdır 
    alti_row_cells = table.rows[5].cells
    alti_row_cells[0].text = 'PLATFORM BAĞLANTI SACI 10mm'
    alti_row_cells[1].text = f'{platform*2}'
    alti_row_cells[2].text = f'{kalinlik1}'
    alti_row_cells[3].text = '3000'
    for i, cell in enumerate(alti_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # yedi satırdaki metinleri yazdır 
    yedi_row_cells = table.rows[6].cells
    yedi_row_cells[0].text = 'KABLO GEÇİŞ BORULARI'
    yedi_row_cells[1].text = f'{platform}'
    yedi_row_cells[2].text = '30mm Boru'
    yedi_row_cells[3].text = f'{uzunluk}'
    for i, cell in enumerate(yedi_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # sekiz satırdaki metinleri yazdır 
    sekiz_row_cells = table.rows[7].cells
    sekiz_row_cells[0].text = 'KABLO BAĞLANTI KUTUSU'
    sekiz_row_cells[1].text = '1'
    sekiz_row_cells[2].text = '230x200'
    sekiz_row_cells[3].text = '400'
    for i, cell in enumerate(sekiz_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # dokuz satırdaki metinleri yazdır 
    dokuz_row_cells = table.rows[8].cells
    dokuz_row_cells[0].text = 'YAN KORKULUK'
    dokuz_row_cells[1].text = f'{yankorkuluk}'
    dokuz_row_cells[2].text = '80x80 Profil'
    dokuz_row_cells[3].text = f'{uzunluk1}'
    for i, cell in enumerate(dokuz_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # on satırdaki metinleri yazdır 
    on_row_cells = table.rows[9].cells
    on_row_cells[0].text = 'MONTAJ KİTİ'
    on_row_cells[1].text = f'{loadcell}'
    on_row_cells[2].text = '250x250'
    on_row_cells[3].text = '210'
    for i, cell in enumerate(on_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # onbir satırdaki metinleri yazdır 
    onbir_row_cells = table.rows[10].cells
    onbir_row_cells[0].text = 'CIVATA & SOMUN'
    onbir_row_cells[1].text = f'{civata}'
    onbir_row_cells[2].text = f'8'
    onbir_row_cells[3].text = f'16'
    for i, cell in enumerate(onbir_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # oniki satırdaki metinleri yazdır 
    oniki_row_cells = table.rows[11].cells
    oniki_row_cells[0].text = 'KAYNAK CİNSİ'
    oniki_row_cells[1].text = 'GAZ ALTI KAYNAK'
    oniki_row_cells[2].text = f''
    oniki_row_cells[3].text = f''
    for i, cell in enumerate(oniki_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
     # onuc satırdaki metinleri yazdır 
    onuc_row_cells = table.rows[12].cells
    onuc_row_cells[0].text = 'BOYA 3 KAT'
    onuc_row_cells[1].text = f'1 KAT ASTAR - 2 KAT ENDÜTRİYEL'
    onuc_row_cells[2].text = f''
    onuc_row_cells[3].text = f''
    for i, cell in enumerate(onuc_row_cells):
        cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
    
    oniki_row_cells[1].merge(oniki_row_cells[3])
    onuc_row_cells[1].merge(onuc_row_cells[3])
    if(usmodel=='B'):
        # ondort satırdaki metinleri yazdır 
        ondort_row_cells = table.rows[13].cells
        ondort_row_cells[0].text = 'HASIR ÇELİK'
        ondort_row_cells[1].text = f'{hasir}m2       5mm'
        ondort_row_cells[2].text = f''
        ondort_row_cells[3].text = f''
        for i, cell in enumerate(ondort_row_cells):
            cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
        # onbes satırdaki metinleri yazdır 
        onbes_row_cells = table.rows[14].cells
        onbes_row_cells[0].text = 'NEVRÜLLÜ DEMİR'
        onbes_row_cells[1].text = f'{platform*5*2}'
        onbes_row_cells[2].text = '16'
        onbes_row_cells[3].text = f'{uzunluk1}'
        for i, cell in enumerate(onbes_row_cells):
            cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
        # onalti satırdaki metinleri yazdır 
        onalti_row_cells = table.rows[15].cells
        onalti_row_cells[0].text = 'C50 BETON'
        onalti_row_cells[1].text = f'{format((uzunluk*((3000-100-uzunluk2)/1000)*(kalinlik1/1000)),".2f")}m3'
        onalti_row_cells[2].text = f''
        onalti_row_cells[3].text = f''
        for i, cell in enumerate(onalti_row_cells):
            cell.paragraphs[0].runs[0].font.color.rgb = mavi  # Mavi renk
        ondort_row_cells[1].merge(ondort_row_cells[3])
        onalti_row_cells[1].merge(onalti_row_cells[3])
    satir5 = doc.add_paragraph(f'\n\n\n\n\n\n\n\n')
    # ******************************************
     # ******************************************
    baslik_metni2 = "ELEKTRONİK AKSAM:" # Başlık metnini ekleyin
    baslik_paragraf2 = doc.add_paragraph() # Başlık paragrafını oluşturun
    baslik_paragraf2.add_run(baslik_metni2).bold = True # Başlık metnini ekleyin ve kalın yapın
    baslik_paragraf2.style = "Heading 2" # Başlık paragrafının stilini ayarlayın
    # ******************************************
    baslik_metni3 = "LOADCELL(YÜK HÜCRESİ) ÖZELLİKLERİ:" # Başlık metnini ekleyin
    baslik_paragraf3 = doc.add_paragraph() # Başlık paragrafını oluşturun
    baslik_paragraf3.add_run(baslik_metni3).bold = True # Başlık metnini ekleyin ve kalın yapın
    baslik_paragraf3.style = "Heading 3" # Başlık paragrafının stilini ayarlayın
    # ******************************************
    satir3 = doc.add_paragraph(f'Ölçüm sistemindeki yük hücreleri paslanmaz çelik ve 30 Ton olup {loadcell} adet kullanılacaktır. Zarar verilecek nominal yük değerinin %150’si kadardır hata oranı % 0.02’dir.  -30 + 60 C Derece arasında çalışmaktadır. Dijital sistemler PTB, OIML onaylıdır 15 tona kadar hava basıncına dayanır. Azami yükte dahi hassasiyetle çalışır. Tüm yük hücreleri aşırı yük korumalıdır. Toz ve neme karşı veya çevre şartlarına karşı korumalıdır.')
    # ******************************************
    baslik_metni4 = "İNDİKATÖR ÖZELLİKLERİ:" # Başlık metnini ekleyin
    baslik_paragraf4 = doc.add_paragraph() # Başlık paragrafını oluşturun
    baslik_paragraf4.add_run(baslik_metni4).bold = True # Başlık metnini ekleyin ve kalın yapın
    baslik_paragraf4.style = "Heading 3" # Başlık paragrafının stilini ayarlayın
    # ******************************************
    satir4 = doc.add_paragraph(f'Dijital gösterge terminali {indikator} tartı ünitesidir günümüzün modern teknolojisi ile elektronik taşıt kantarları için tasarlanmıştır. Sahip olduğu programla kullanıcının tüm beklentilerini karşılar. Hafızada depolanan bilgiler istenilen düzeyde ekrandan veya yazıcıdan rapor halinde alınabilir. Uluslararası OIML IR 76 standartlarında ve Avrupa birliği 90/384/EEC tip onayına sahiptir.')
    # ******************************************
    # Metin ve linki tanımla
    # metin = f'{linkB6}'
    # satir4 = doc.add_paragraph('İnşaat Projesi :')
    # satir41 = add_hyperlink(satir4, 'https://drive.google.com/file/d/1el2eAMfcprVtToEwUxMSUF60ZRvoiHOm/view?usp=sharing', 'B6 inşaat projesi.pdf', '1E90FF', False)
    # run = satir4.add_run(satir41)
    # Metin ve linki tanımla

   
    # ***************AYAR**********************
    # hizalama
    satir1.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir11.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    satir4.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    # satir5.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


    # Belgeyi kaydet
    word_dizini = 'C:/Users/yenic/Desktop/Teklif_Django_Template/otomatik/word/'
    doc.save(f'{word_dizini}/{dosya2}.docx')
    print(f'{dosya2} word belgesi olusturuldu')



# towordtecnical()